#!/usr/bin/env python3
"""
RyderVocabFormatter - 小朋友英语词汇文档格式化工具
Version: 1.1.0 - 修复章节识别和混合行解析
Author: Nova
"""

import argparse
import sys
import os
import re
import glob
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn

# 默认配置
DEFAULT_CONFIG = {
    'highlight_color': WD_COLOR_INDEX.YELLOW,
    'line_spacing': 1.1,
    'font_english': 'Times New Roman',
    'font_chinese': '宋体',
    'font_size': 11,
    'title_font_size': 14,
    'section_font_size': 12,
}

def set_run_font(run, text, config):
    """智能设置字体"""
    english_chars = len(re.findall(r'[a-zA-Z]', text))
    chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
    
    if english_chars > chinese_chars:
        run.font.name = config['font_english']
    else:
        run.font.name = config['font_chinese']
    run._element.rPr.rFonts.set(qn('w:eastAsia'), config['font_chinese'])
    run.font.size = Pt(config['font_size'])

def split_mixed_line(text):
    """
    分离混合行
    例如："1. 单词整理business /ˈbɪznəs/ 名词..." 
    分离为：["1. 单词整理", "business /ˈbɪznəs/ 名词..."]
    
    例如："dragon army 龙军文中例句：..."
    分离为：["dragon army 龙军", "文中例句：..."]
    """
    lines = []
    
    # 检查是否是章节标题 + 单词混合
    # 模式：章节标题 + 英文单词 + 音标
    section_word_pattern = r'^(\d+[\.。]\s*\S+整理)\s*([a-zA-Z]+\s*/[^/]+/.*)$'
    match = re.match(section_word_pattern, text)
    if match:
        lines.append(match.group(1))
        lines.append(match.group(2))
        return lines
    
    # 检查是否是短语 + 例句混合（没有换行的情况）
    # 模式：英文短语 + 中文翻译 + 文中例句
    phrase_example_pattern = r'^([a-zA-Z\s]+\s+[\u4e00-\u9fff]+)(文中例句.*)$'
    match = re.match(phrase_example_pattern, text)
    if match:
        lines.append(match.group(1).strip())
        lines.append(match.group(2).strip())
        return lines
    
    # 检查是否包含"翻译"但没有分开
    translation_pattern = r'^(.*)(翻译[：:].*)$'
    match = re.match(translation_pattern, text)
    if match and '文中例句' not in match.group(1):
        lines.append(match.group(1).strip())
        lines.append(match.group(2).strip())
        return lines
    
    lines.append(text)
    return lines

def identify_section_title(text):
    """识别章节标题，支持多种格式"""
    # 清理文本
    clean = re.sub(r'\s+', ' ', text.strip())
    
    # 单词整理 - 支持多种格式（更宽松的匹配）
    if re.search(r'\d+.*单词.*整理', clean):
        return 'words'
    
    # 短语整理
    if re.search(r'\d+.*短语.*整理', clean):
        return 'phrases'
    
    # 常考重点词汇
    if re.search(r'\d+.*常考.*词汇', clean):
        return 'key_points'
    
    # 重点单词词性转换
    if '重点单词' in clean and '词性' in clean:
        return 'key_points'
    
    return None

def identify_item_type(text, section):
    """识别内容类型"""
    text_clean = re.sub(r'^-\s*', '', text)
    
    # 检测是否是翻译行（优先检查，因为翻译通常包含"翻译"关键词）
    if text.startswith('翻译：') or text.startswith('翻译'):
        return 'translation'
    
    if section == 'words':
        # 单词特征：音标 /.../
        if re.search(r'/[ɑæɛɪɒʊʌəɜɪʊɔɑɜɪʊəɜʃθðŋʒʤtʃdʒkw]+/', text_clean):
            return 'word'
        # 或词性标注
        if any(pos in text_clean for pos in ['名词', '动词', '形容词', '副词', '数词', '介词']):
            return 'word'
        # 检测例句
        if '文中例句' in text:
            return 'example'
        if text.startswith('例句：'):
            return 'example'
    
    elif section == 'phrases':
        # 短语特征：英文+中文，无音标（优先于例句检查）
        if re.search(r'[a-zA-Z]', text_clean) and re.search(r'[\u4e00-\u9fff]', text_clean):
            if not re.search(r'/[ɑæɛɪɒʊʌəɜɪʊɔɑɜɪʊəɜʃθðŋʒʤtʃdʒkw]+/', text_clean):
                # 即使包含"文中例句"，只要前面有英文+中文，就认为是短语
                return 'phrase'
        # 然后是例句
        if '文中例句' in text:
            return 'example'
    
    elif section == 'key_points':
        if re.match(r'^-?\s*(动词|名词|形容词|副词|介词)', text_clean):
            return 'pos'
        if '例句' in text and '：' in text:
            return 'example'
        if '→' in text or '词性转' in text_clean:
            return 'conversion'
        if re.match(r'^[a-zA-Z\s]+$', text_clean.strip()):
            return 'key_point'
    
    return 'other'

def parse_document(input_file):
    """解析文档，处理混合行"""
    doc = Document(input_file)
    
    sections_data = {
        'words': {'title': '1. 单词整理', 'items': []},
        'phrases': {'title': '2. 短语整理', 'items': []},
        'key_points': {'title': '3. 常考重点词汇不同词性及例句', 'items': []},
    }
    
    current_section = None
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # 尝试识别章节标题
        section_type = identify_section_title(text)
        if section_type:
            current_section = section_type
            # 如果标题行还包含内容，分离出来
            lines = split_mixed_line(text)
            if len(lines) > 1:
                # 第一行是标题，后续行是内容
                for line in lines[1:]:
                    if current_section:
                        item_type = identify_item_type(line, current_section)
                        sections_data[current_section]['items'].append({
                            'text': line,
                            'type': item_type
                        })
            continue
        
        # 处理普通行，检查是否是混合行
        lines = split_mixed_line(text)
        for line in lines:
            if current_section:
                item_type = identify_item_type(line, current_section)
                sections_data[current_section]['items'].append({
                    'text': line,
                    'type': item_type
                })
    
    return sections_data

def merge_examples(items):
    """合并重复例句"""
    example_map = {}
    for i, item in enumerate(items):
        if item['type'] == 'example':
            text = item['text']
            if text not in example_map:
                example_map[text] = []
            example_map[text].append(i)
    
    to_remove = set()
    for example_text, indices in example_map.items():
        if len(indices) > 1:
            for idx in indices[:-1]:
                to_remove.add(idx)
                if idx + 1 < len(items) and items[idx + 1]['type'] == 'translation':
                    to_remove.add(idx + 1)
    
    return [item for i, item in enumerate(items) if i not in to_remove]

def process_items(doc, items, item_type, config):
    """处理单词或短语章节"""
    counter = 0
    
    for item in items:
        text = re.sub(r'^-\s*', '', item['text'])
        para = doc.add_paragraph()
        para.paragraph_format.line_spacing = config['line_spacing']
        
        if item['type'] == item_type:
            counter += 1
            text = re.sub(r'^\d+[\.。]\s*', '', text)  # 移除原有编号
            text = f"{counter}. {text}"
            run = para.add_run(text)
            run.font.highlight_color = config['highlight_color']
        elif item['type'] in ['example', 'translation']:
            run = para.add_run(text)
        else:
            run = para.add_run(text)
        
        set_run_font(run, text, config)
    
    return counter

def process_key_points(doc, items, config):
    """处理常考重点词汇章节"""
    counter = 0
    
    for item in items:
        text = re.sub(r'^-\s*', '', item['text'])
        para = doc.add_paragraph()
        para.paragraph_format.line_spacing = config['line_spacing']
        
        if item['type'] == 'key_point':
            counter += 1
            text = f"{counter}. {text}"
            run = para.add_run(text)
            run.font.highlight_color = config['highlight_color']
        elif item['type'] == 'pos':
            text = f"- {text}"
            run = para.add_run(text)
        else:
            run = para.add_run(text)
        
        set_run_font(run, text, config)
    
    return counter

def process_document(input_file, output_file, config=None):
    """处理文档"""
    if config is None:
        config = DEFAULT_CONFIG.copy()
    
    # 解析文档
    sections_data = parse_document(input_file)
    
    # 创建新文档
    new_doc = Document()
    
    # 复制页面设置
    old_doc = Document(input_file)
    for attr in ['page_height', 'page_width', 'left_margin', 'right_margin', 'top_margin', 'bottom_margin']:
        setattr(new_doc.sections[0], attr, getattr(old_doc.sections[0], attr))
    
    # 添加文档标题
    title_para = new_doc.add_paragraph()
    title_para.paragraph_format.line_spacing = config['line_spacing']
    run = title_para.add_run("Week2-3 3.7")
    run.font.size = Pt(config['title_font_size'])
    run.font.bold = True
    set_run_font(run, "Week2-3 3.7", config)
    
    # 处理每个章节
    stats = {}
    for section_key in ['words', 'phrases', 'key_points']:
        section = sections_data[section_key]
        if not section['items']:
            continue
        
        # 章节标题
        title_para = new_doc.add_paragraph()
        title_para.paragraph_format.line_spacing = config['line_spacing']
        run = title_para.add_run(section['title'])
        run.font.size = Pt(config['section_font_size'])
        run.font.bold = True
        set_run_font(run, section['title'], config)
        
        # 合并例句并处理
        items = merge_examples(section['items'])
        
        if section_key == 'phrases':
            count = process_items(new_doc, items, 'phrase', config)
            stats['phrases'] = count
        elif section_key == 'words':
            count = process_items(new_doc, items, 'word', config)
            stats['words'] = count
        elif section_key == 'key_points':
            count = process_key_points(new_doc, items, config)
            stats['key_points'] = count
    
    new_doc.save(output_file)
    return stats

def main():
    parser = argparse.ArgumentParser(
        description='RyderVocabFormatter - 小朋友英语词汇文档格式化工具 v1.1.0',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  %(prog)s -i vocab.docx                    # 处理单个文件
  %(prog)s -i vocab.docx -o output.docx     # 指定输出文件
  %(prog)s --batch "week*.docx"             # 批量处理
  %(prog)s -i vocab.docx --dry-run          # 预览模式
        """
    )
    
    parser.add_argument('-i', '--input', required=True, help='输入文件或通配符模式')
    parser.add_argument('-o', '--output', help='输出文件（批量处理时指定目录）')
    parser.add_argument('--batch', action='store_true', help='批量处理模式')
    parser.add_argument('--dry-run', action='store_true', help='预览模式（不保存文件）')
    parser.add_argument('--line-spacing', type=float, default=1.1, help='行间距（默认1.1）')
    parser.add_argument('--font-en', default='Times New Roman', help='英文字体')
    parser.add_argument('--font-cn', default='宋体', help='中文字体')
    
    args = parser.parse_args()
    
    # 更新配置
    config = DEFAULT_CONFIG.copy()
    config['line_spacing'] = args.line_spacing
    config['font_english'] = args.font_en
    config['font_chinese'] = args.font_cn
    
    if args.batch:
        # 批量处理
        pattern = args.input
        files = glob.glob(pattern)
        if not files:
            print(f"❌ 未找到匹配的文件: {pattern}")
            return 1
        
        output_dir = args.output or '.'
        os.makedirs(output_dir, exist_ok=True)
        
        print(f"📝 批量处理 {len(files)} 个文件...\n")
        
        for file in sorted(files):
            basename = os.path.basename(file)
            output_file = os.path.join(output_dir, f"formatted_{basename}")
            
            if args.dry_run:
                print(f"📄 [预览] {basename} → {output_file}")
            else:
                try:
                    stats = process_document(file, output_file, config)
                    print(f"✅ {basename}")
                    print(f"   单词: {stats.get('words', 0)} | 短语: {stats.get('phrases', 0)} | 重点: {stats.get('key_points', 0)}")
                except Exception as e:
                    print(f"❌ {basename} - 错误: {e}")
        
        print(f"\n🎉 批量处理完成！输出目录: {output_dir}")
    
    else:
        # 单文件处理
        input_file = args.input
        if not os.path.exists(input_file):
            print(f"❌ 文件不存在: {input_file}")
            return 1
        
        output_file = args.output or input_file.replace('.docx', '_formatted.docx')
        
        if args.dry_run:
            print(f"📄 [预览模式] 将处理: {input_file}")
            print(f"   输出: {output_file}")
            return 0
        
        try:
            stats = process_document(input_file, output_file, config)
            print(f"✅ 处理完成: {output_file}")
            print(f"   单词: {stats.get('words', 0)} 个")
            print(f"   短语: {stats.get('phrases', 0)} 个")
            print(f"   重点词汇: {stats.get('key_points', 0)} 个")
        except Exception as e:
            print(f"❌ 处理失败: {e}")
            import traceback
            traceback.print_exc()
            return 1
    
    return 0

if __name__ == '__main__':
    sys.exit(main())
